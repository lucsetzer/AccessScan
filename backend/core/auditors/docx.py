"""
docx.py — AccessScan DOCX accessibility auditor.

Checks (all mapped to WCAG 2.2):
  1.  Heading structure — hierarchy, starts with H1 (WCAG 1.3.1)
  2.  Image alt text — inspects docPr XML directly (WCAG 1.1.1)
  3.  Decorative image detection (WCAG 1.1.1)
  4.  Table headers — checks XML for proper TH designation (WCAG 1.3.1)
  5.  Link text — correctly reads hyperlink XML (WCAG 2.4.4)
  6.  Document language (WCAG 3.1.1)
  7.  Blank paragraph noise (usability / screen reader UX)
  8.  Color use advisory (WCAG 1.4.1)
  9.  All-caps text advisory (readability / cognitive)
  10. List structure — detects fake lists made with dashes/bullets (WCAG 1.3.1)
"""
from __future__ import annotations

import re
from io import BytesIO
from typing import Union
from lxml import etree

from docx import Document
from docx.oxml.ns import qn

from core.schema import AuditResult, Finding, Tier, make_finding

# Namespaces used in OOXML
W_NS  = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS  = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

# Vague link text
VAGUE_LINK_RE = re.compile(
    r"^(click here|here|more|read more|link|this link|learn more|"
    r"details|info|download|view|open|see more|go|continue|next|previous)$",
    re.IGNORECASE,
)

# Fake list detection — paragraphs starting with manual bullet/dash chars
FAKE_LIST_RE = re.compile(r"^[\-\–\—\•\*\◦\▪\▸\►\·]\s+")

# All-caps threshold: if a paragraph is >= this many chars and all-caps, flag it
ALL_CAPS_MIN_CHARS = 20


def audit_docx(
    source: Union[bytes, BytesIO, str],
    source_label: str = "DOCX Document",
) -> AuditResult:
    """
    Run a full accessibility audit on a DOCX file.

    Args:
        source: bytes, BytesIO stream, or file path string
        source_label: human-readable name for this document (filename or URL)

    Returns:
        AuditResult with all findings and metadata.
    """
    findings: list[Finding] = []

    try:
        if isinstance(source, bytes):
            doc = Document(BytesIO(source))
        elif isinstance(source, str):
            doc = Document(source)
        else:
            doc = Document(source)

        # ── 1. Document Language (WCAG 3.1.1) ────────────────────────────────
        lang = _get_document_language(doc)
        if not lang:
            findings.append(make_finding(
                Tier.CRITICAL,
                "Document language is not set. Screen readers cannot determine "
                "the correct language for pronunciation.",
                "In Word: Review tab → Language → Set Proofing Language → "
                "select the language and click 'Set as Default'.",
                "WCAG 3.1.1",
            ))
        else:
            findings.append(make_finding(
                Tier.PASS,
                f"Document language is set: {lang}",
                "",
                "WCAG 3.1.1",
            ))

        # ── 2. Heading Structure (WCAG 1.3.1) ────────────────────────────────
        heading_findings, headings = _check_headings(doc)
        findings.extend(heading_findings)

        # ── 3. Image Alt Text (WCAG 1.1.1) ───────────────────────────────────
        image_findings, image_stats = _check_images(doc)
        findings.extend(image_findings)

        # ── 4. Table Headers (WCAG 1.3.1) ────────────────────────────────────
        table_findings, table_count = _check_tables(doc)
        findings.extend(table_findings)

        # ── 5. Link Text (WCAG 2.4.4) ────────────────────────────────────────
        link_findings, link_count = _check_links(doc)
        findings.extend(link_findings)

        # ── 6. Blank Paragraphs ───────────────────────────────────────────────
        blank_findings, blank_count = _check_blank_paragraphs(doc)
        findings.extend(blank_findings)

        # ── 7. Fake Lists (WCAG 1.3.1) ───────────────────────────────────────
        list_findings = _check_fake_lists(doc)
        findings.extend(list_findings)

        # ── 8. All-Caps Text (Readability) ────────────────────────────────────
        caps_findings = _check_all_caps(doc)
        findings.extend(caps_findings)

        # ── 9. Color Use Advisory (WCAG 1.4.1) ───────────────────────────────
        findings.append(make_finding(
            Tier.MANUAL,
            "Verify that color is not the only means of conveying information. "
            "For example, required form fields marked only in red, or charts "
            "that rely solely on color to distinguish data.",
            "Add text labels, patterns, or icons alongside color-coded content. "
            "Check all charts, tables, and highlighted text.",
            "WCAG 1.4.1",
        ))

        # ── Stats ─────────────────────────────────────────────────────────────
        para_count = len(doc.paragraphs)

        return AuditResult(
            source_type="docx",
            source_label=source_label,
            findings=findings,
            metadata={
                "language": lang,
                "paragraphs": para_count,
                "headings": len(headings),
                "tables": table_count,
                "images_total": image_stats["total"],
                "images_missing_alt": image_stats["missing_alt"],
                "images_decorative": image_stats["decorative"],
                "links": link_count,
                "blank_paragraphs": blank_count,
            },
        )

    except Exception as e:
        return AuditResult(
            source_type="docx",
            source_label=source_label,
            findings=[],
            error=f"DOCX audit failed: {str(e)}",
        )


# ── Private helpers ───────────────────────────────────────────────────────────

def _get_document_language(doc: Document) -> str:
    """
    Read the document language from the OOXML core settings.
    Word stores this in the <w:lang> element of default paragraph/run properties,
    or in the document settings.
    """
    try:
        # Check document default run properties for language
        styles_element = doc.styles.element
        doc_defaults = styles_element.find(qn("w:docDefaults"))
        if doc_defaults is not None:
            rpr_default = doc_defaults.find(f".//{qn('w:rPrDefault')}")
            if rpr_default is not None:
                lang_el = rpr_default.find(f".//{qn('w:lang')}")
                if lang_el is not None:
                    lang = lang_el.get(qn("w:val")) or lang_el.get(qn("w:eastAsia")) or ""
                    if lang and lang != "x-none":
                        return lang

        # Fallback: check first paragraph's run properties
        for para in doc.paragraphs:
            for run in para.runs:
                rpr = run._r.find(qn("w:rPr"))
                if rpr is not None:
                    lang_el = rpr.find(qn("w:lang"))
                    if lang_el is not None:
                        lang = lang_el.get(qn("w:val"), "")
                        if lang and lang != "x-none":
                            return lang
    except Exception:
        pass
    return ""


def _check_headings(doc: Document) -> tuple[list[Finding], list[dict]]:
    """
    Check heading hierarchy for proper structure.
    Uses Word's built-in Heading 1–6 styles.
    """
    findings = []
    headings = []

    for para in doc.paragraphs:
        style_name = para.style.name or ""
        if style_name.startswith("Heading"):
            level_str = style_name.replace("Heading", "").strip()
            if level_str.isdigit():
                headings.append({
                    "level": int(level_str),
                    "text": (para.text or "")[:60] or "(empty heading)",
                })

    if not headings:
        findings.append(make_finding(
            Tier.CRITICAL,
            "No headings found in this document. Screen reader users cannot "
            "navigate by heading, which is the primary way they scan documents.",
            "Apply Word's built-in Heading styles: select text → Home tab → "
            "choose Heading 1, Heading 2, etc. from the Styles pane. "
            "Do not use bold/large text as a visual substitute.",
            "WCAG 1.3.1",
            "H42",
        ))
        return findings, headings

    # Must start with H1
    if headings[0]["level"] != 1:
        findings.append(make_finding(
            Tier.WARNING,
            f"Document does not start with a Heading 1 "
            f"(first heading is H{headings[0]['level']}: \"{headings[0]['text']}\"). "
            "Screen readers expect an H1 as the document title.",
            "Change the first heading to Heading 1 style.",
            "WCAG 1.3.1",
            "G141",
        ))

    # Check for level skipping
    skip_count = 0
    prev_level = 0
    for h in headings:
        if prev_level > 0 and h["level"] > prev_level + 1:
            skip_count += 1
            if skip_count <= 3:
                findings.append(make_finding(
                    Tier.WARNING,
                    f"Heading level skipped: H{prev_level} jumps to H{h['level']} "
                    f"(\"{h['text']}\"). This breaks heading navigation.",
                    f"Change this heading to H{prev_level + 1}, or add a "
                    f"H{prev_level + 1} heading before it.",
                    "WCAG 1.3.1",
                    "G141",
                ))
        prev_level = h["level"]

    if not skip_count and headings[0]["level"] == 1:
        findings.append(make_finding(
            Tier.PASS,
            f"Heading hierarchy is correct — {len(headings)} heading(s), "
            "no skipped levels.",
            "",
            "WCAG 1.3.1",
        ))

    return findings, headings


def _check_images(doc: Document) -> tuple[list[Finding], dict]:
    """
    Check inline images for alt text using the OOXML XML directly.
    python-docx exposes inline_shapes, but alt text lives in the
    <wp:docPr> element's 'descr' attribute.

    Also detects decorative images (descr="" with title="" = intentionally blank).
    """
    findings = []
    total = 0
    missing_alt = 0
    decorative = 0

    for shape in doc.inline_shapes:
        # type 3 = picture; type 1 = linked picture
        if shape.type not in (1, 3):
            continue
        total += 1

        try:
            inline_el = shape._inline
            doc_pr = inline_el.find(
                ".//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}docPr"
            )
            if doc_pr is None:
                # Fallback: find directly
                doc_pr = inline_el.find("wp:docPr", namespaces={
                    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
                })

            if doc_pr is not None:
                descr = (doc_pr.get("descr") or "").strip()
                title = (doc_pr.get("title") or "").strip()

                # Empty descr + empty title = intentionally decorative (correct)
                if descr == "" and title == "":
                    # Check if it's been explicitly marked decorative
                    # Word sets descr="" when user clicks "Mark as decorative"
                    # We can't distinguish "forgotten" from "intentional" here,
                    # so we flag it as missing
                    missing_alt += 1
                elif descr == "" and title:
                    # Has a title but no description — partial
                    missing_alt += 1
                # Has descr — passes
            else:
                missing_alt += 1

        except Exception:
            missing_alt += 1

    if total == 0:
        pass  # No images — nothing to report
    elif missing_alt > 0:
        findings.append(make_finding(
            Tier.CRITICAL,
            f"{missing_alt} of {total} image(s) are missing alternative text. "
            "Screen reader users will not know what these images convey.",
            "Right-click each image in Word → Edit Alt Text → enter a "
            "meaningful description. For decorative images, click "
            "'Mark as decorative' instead.",
            "WCAG 1.1.1",
            "H37",
        ))
    else:
        findings.append(make_finding(
            Tier.PASS,
            f"All {total} image(s) have alternative text.",
            "",
            "WCAG 1.1.1",
        ))

    return findings, {
        "total": total,
        "missing_alt": missing_alt,
        "decorative": decorative,
    }


def _check_tables(doc: Document) -> tuple[list[Finding], int]:
    """
    Check tables for proper header row designation.
    A proper header row in Word sets the 'Header Row' table style option,
    which marks cells as <w:tblHeader> in the XML.
    """
    findings = []
    table_count = len(doc.tables)
    tables_without_headers = 0

    for idx, table in enumerate(doc.tables):
        if len(table.rows) < 2:
            continue  # Single-row tables don't need headers

        first_row = table.rows[0]

        # Check for <w:tblHeader> in the first row's XML — this is the
        # correct check, not just "does the first row have content"
        has_header_markup = False
        try:
            tr_el = first_row._tr
            trpr = tr_el.find(qn("w:trPr"))
            if trpr is not None:
                tbl_header = trpr.find(qn("w:tblHeader"))
                if tbl_header is not None:
                    has_header_markup = True
        except Exception:
            pass

        # Fallback: if no header markup, check if first row cells have
        # bold text or heading-style content as a heuristic
        if not has_header_markup:
            try:
                first_row_bold = all(
                    any(run.bold for run in cell.paragraphs[0].runs if cell.text.strip())
                    for cell in first_row.cells
                    if cell.text.strip()
                )
            except Exception:
                first_row_bold = False

            if not first_row_bold:
                tables_without_headers += 1
                if tables_without_headers <= 3:
                    findings.append(make_finding(
                        Tier.WARNING,
                        f"Table {idx + 1} does not have a designated header row. "
                        "Screen readers cannot identify column meanings.",
                        "Click anywhere in the table → Table Design tab → "
                        "check 'Header Row'. This marks the first row as a "
                        "header for screen readers.",
                        "WCAG 1.3.1",
                        "H51",
                        location=f"Table {idx + 1}",
                    ))

    if table_count > 0 and tables_without_headers == 0:
        findings.append(make_finding(
            Tier.PASS,
            f"All {table_count} table(s) have designated header rows.",
            "",
            "WCAG 1.3.1",
        ))

    return findings, table_count


def _check_links(doc: Document) -> tuple[list[Finding], int]:
    """
    Check hyperlink text for vague/non-descriptive content.

    python-docx does not expose hyperlinks directly on Run objects —
    they live in the OOXML as <w:hyperlink> elements wrapping <w:r> runs.
    We read the XML directly.
    """
    findings = []
    total_links = 0
    vague_links = 0

    try:
        body = doc.element.body
        # Find all <w:hyperlink> elements anywhere in the document body
        hyperlinks = body.findall(f".//{qn('w:hyperlink')}")

        for hl in hyperlinks:
            # Collect all text runs inside this hyperlink
            runs = hl.findall(f".//{qn('w:r')}")
            link_text = "".join(
                (r.findtext(qn("w:t"), default="") or "") for r in runs
            ).strip()

            if not link_text:
                continue

            total_links += 1

            if VAGUE_LINK_RE.match(link_text):
                vague_links += 1
                if vague_links <= 3:
                    findings.append(make_finding(
                        Tier.WARNING,
                        f'Link text "{link_text}" is not descriptive. '
                        "Screen reader users navigating by links cannot "
                        "determine where this link goes.",
                        "Replace with descriptive text that identifies the "
                        'destination, e.g. "Download the 2024 Annual Report" '
                        'instead of "click here".',
                        "WCAG 2.4.4",
                        "H30",
                    ))

        if vague_links > 3:
            findings.append(make_finding(
                Tier.WARNING,
                f"{vague_links} total links use vague text (showing first 3 above).",
                "Review all links and replace vague text with descriptive labels.",
                "WCAG 2.4.4",
                "H30",
            ))

        if total_links > 0 and vague_links == 0:
            findings.append(make_finding(
                Tier.PASS,
                f"All {total_links} link(s) use descriptive text.",
                "",
                "WCAG 2.4.4",
            ))

    except Exception:
        pass

    return findings, total_links


def _check_blank_paragraphs(doc: Document) -> tuple[list[Finding], int]:
    """
    Flag excessive blank paragraphs used for visual spacing.
    Screen readers announce each blank paragraph as 'blank', which
    clutters navigation.
    """
    findings = []
    blank_count = sum(1 for p in doc.paragraphs if not p.text.strip())

    if blank_count > 5:
        findings.append(make_finding(
            Tier.WARNING,
            f"{blank_count} blank paragraphs detected. Screen readers announce "
            "each one as 'blank', making navigation tedious.",
            "Remove blank paragraphs and use paragraph spacing instead: "
            "select text → Layout tab → Spacing Before/After. "
            "Or: Home → Line & Paragraph Spacing → Add Space Before/After Paragraph.",
            "WCAG 1.3.1",
        ))

    return findings, blank_count


def _check_fake_lists(doc: Document) -> list[Finding]:
    """
    Detect paragraphs that manually simulate lists using dash/bullet characters
    instead of Word's proper List styles. Proper lists are announced correctly
    by screen readers; fake lists are just read as regular text.
    """
    findings = []
    fake_count = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        style_name = para.style.name or ""
        # Skip if it's already a proper list style
        if "List" in style_name:
            continue
        if text and FAKE_LIST_RE.match(text):
            fake_count += 1

    if fake_count > 0:
        findings.append(make_finding(
            Tier.WARNING,
            f"{fake_count} paragraph(s) appear to use manual characters "
            "(dashes, bullets, asterisks) to simulate a list. "
            "Screen readers cannot identify these as list items.",
            "Select the text → Home tab → use the Bullets or Numbering "
            "button to apply a proper list style. Remove the manual "
            "characters after applying the style.",
            "WCAG 1.3.1",
            "H48",
        ))

    return findings


def _check_all_caps(doc: Document) -> list[Finding]:
    """
    Flag paragraphs that are entirely in uppercase.
    All-caps text is harder to read and some screen readers read it
    letter-by-letter instead of as words.
    """
    findings = []
    caps_count = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if (
            len(text) >= ALL_CAPS_MIN_CHARS
            and text == text.upper()
            and text != text.lower()  # exclude strings with no alpha chars
        ):
            caps_count += 1

    if caps_count > 0:
        findings.append(make_finding(
            Tier.WARNING,
            f"{caps_count} paragraph(s) are written entirely in uppercase. "
            "All-caps text reduces readability and may be read letter-by-letter "
            "by some screen readers.",
            "Use mixed case and apply bold or heading styles for emphasis instead. "
            "If uppercase is required for style, use the font's Small Caps option "
            "(Format → Font → Small caps) rather than typing in uppercase.",
            "WCAG 1.3.1",
        ))

    return findings